"""
generate_docx.py
Membuat template DOCX Tugas Akhir ITS (Teknik Informatika FTEIC)
sesuai format buku TA dari TA-Ricardo.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─────────────────────────────────────────────
# DATA METADATA — sesuaikan di sini
# ─────────────────────────────────────────────
META = {
    "kode_mk": "EF234801",
    "judul_id": "Perancangan Sistem Rekomendasi Author dengan Perlindungan Privasi Query Menggunakan Homomorphic Inference",
    "judul_en": "Design of an Author Recommendation System with Query Privacy Protection Using Homomorphic Inference",
    "nama": "Ricardo Supriyanto",
    "nrp": "5025221218",
    "pembimbing": {"nama": "Moch. Nafkhan Alzamzami, S.T., M.T.", "nip": "199911222024061001"},
    "ko_pembimbing": [
        {"nama": "Ilham Gurat Adillion, S.Kom., M.Eng.", "nip": "1995202411009"},
        {"nama": "Prof. Ir. Ary Mazharuddin Shiddiqi, S.Kom., M.Comp.Sc., Ph.D.", "nip": "198106202005011003"},
    ],
    "penguji": [
        {"nama": "<Nama dan gelar penguji>"},
        {"nama": "<Nama dan gelar penguji>"},
    ],
    "prodi_id": "Program Studi S-1 Teknik Informatika",
    "prodi_en": "Undergraduate Study Program of Informatics",
    "dept_id": "Departemen Teknik Informatika",
    "dept_en": "Department of Informatics",
    "fak_id": "Fakultas Teknologi Elektro dan Informatika Cerdas",
    "fak_en": "Faculty of Intelligent Electrical and Informatics Technology",
    "institusi": "Institut Teknologi Sepuluh Nopember",
    "kota": "Surabaya",
    "tahun": "2026",
    "bulan": "Januari",
    "bulan_en": "January",
    "kata_kunci_id": "sistem rekomendasi author, privasi query, homomorphic inference, fully homomorphic encryption, concrete",
    "kata_kunci_en": "author recommendation system, query privacy, homomorphic inference, fully homomorphic encryption, concrete",
}


# ─────────────────────────────────────────────
# HELPER FUNCTIONS
# ─────────────────────────────────────────────

def set_margins(doc, top=3, bottom=3, left=4, right=3):
    """Set halaman A4 dengan margin ITS (cm)."""
    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)


def add_styles(doc):
    """Buat style paragraph kustom ITS."""
    styles = doc.styles

    def _ensure_style(name, base="Normal"):
        if name in [s.name for s in styles]:
            return styles[name]
        s = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        if base in [st.name for st in styles]:
            s.base_style = styles[base]
        return s

    # ── Normal body ──────────────────────────────────
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = Pt(18)           # 1.5 spasi (12pt × 1.5)
    pf.space_before = Pt(0)
    pf.space_after  = Pt(6)
    pf.first_line_indent = Cm(1.25)

    # ── Heading 1 (BAB) ──────────────────────────────
    h1 = styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1pf = h1.paragraph_format
    h1pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1pf.space_before = Pt(24)
    h1pf.space_after  = Pt(12)
    h1pf.page_break_before = True
    h1pf.first_line_indent = Cm(0)
    h1pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    h1pf.line_spacing = Pt(18)

    # ── Heading 2 (sub-bab) ──────────────────────────
    h2 = styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2pf = h2.paragraph_format
    h2pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2pf.space_before = Pt(12)
    h2pf.space_after  = Pt(6)
    h2pf.first_line_indent = Cm(0)
    h2pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    h2pf.line_spacing = Pt(18)

    # ── Heading 3 (sub-sub-bab) ───────────────────────
    h3 = styles["Heading 3"]
    h3.font.name = "Times New Roman"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3pf = h3.paragraph_format
    h3pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3pf.space_before = Pt(6)
    h3pf.space_after  = Pt(3)
    h3pf.first_line_indent = Cm(0)
    h3pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    h3pf.line_spacing = Pt(18)

    # ── Caption gambar / tabel ───────────────────────
    cap = _ensure_style("Caption ITS", "Normal")
    cap.font.name = "Times New Roman"
    cap.font.size = Pt(11)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after  = Pt(6)

    # ── Kode sumber ──────────────────────────────────
    code = _ensure_style("Kode Sumber", "Normal")
    code.font.name = "Courier New"
    code.font.size = Pt(10)
    code.paragraph_format.first_line_indent = Cm(0)
    code.paragraph_format.left_indent  = Cm(1)
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after  = Pt(0)
    code.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # ── Judul halaman khusus (ABSTRAK, dll) ──────────
    title_sp = _ensure_style("Judul Halaman", "Normal")
    title_sp.font.name = "Times New Roman"
    title_sp.font.size = Pt(14)
    title_sp.font.bold = True
    title_sp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_sp.paragraph_format.first_line_indent = Cm(0)
    title_sp.paragraph_format.space_before = Pt(0)
    title_sp.paragraph_format.space_after  = Pt(12)

    # ── Teks tengah (cover) ──────────────────────────
    center = _ensure_style("Center Text", "Normal")
    center.font.name = "Times New Roman"
    center.font.size = Pt(12)
    center.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    center.paragraph_format.first_line_indent = Cm(0)
    center.paragraph_format.space_before = Pt(0)
    center.paragraph_format.space_after  = Pt(0)


def add_page_break(doc):
    doc.add_page_break()


def add_blank_intentional(doc):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.style = doc.styles["Center Text"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(300)
    run = p.add_run("Halaman ini sengaja dikosongkan.")
    run.italic = True


def center_para(doc, text, bold=False, size=12, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.style = doc.styles["Center Text"]
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def body_para(doc, text="", bold=False, indent=True):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    if not indent:
        p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.bold = bold
    return p


def add_2col_table(doc, rows_data, widths=(5, 0.5, 10)):
    """Tabel 3 kolom: label | : | value (tanpa border)."""
    tbl = doc.add_table(rows=len(rows_data), cols=3)
    tbl.style = "Table Grid"
    # Hapus semua border
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = OxmlElement(f"w:{edge}")
                border.set(qn("w:val"), "none")
                tcBorders.append(border)
            tcPr.append(tcBorders)
    # Set lebar kolom
    for i, row_data in enumerate(rows_data):
        cells = tbl.rows[i].cells
        for j, (val, w) in enumerate(zip(row_data, widths)):
            cells[j].width = Cm(w)
            p = cells[j].paragraphs[0]
            p.style = doc.styles["Normal"]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = Pt(18)
            run = p.add_run(str(val))
            if j == 0:
                run.bold = False
    return tbl


# ─────────────────────────────────────────────
# BAGIAN-BAGIAN DOKUMEN
# ─────────────────────────────────────────────

def add_cover(doc, lang="id"):
    add_page_break(doc) if doc.paragraphs else None

    judul = META["judul_id"] if lang == "id" else META["judul_en"]
    tipe  = f"TUGAS AKHIR — {META['kode_mk']}" if lang == "id" else f"FINAL PROJECT — {META['kode_mk']}"
    prodi = META["prodi_id"] if lang == "id" else META["prodi_en"]
    dept  = META["dept_id"]  if lang == "id" else META["dept_en"]
    fak   = META["fak_id"]   if lang == "id" else META["fak_en"]

    center_para(doc, tipe, bold=True, size=14, space_before=0, space_after=36)
    center_para(doc, judul.upper(), bold=True, size=14, space_after=48)
    center_para(doc, META["nama"], size=12, space_after=0)
    center_para(doc, f"NRP {META['nrp']}", size=12, space_after=36)

    supervisors = [
        ("Dosen Pembimbing" if lang == "id" else "Advisor", META["pembimbing"]),
        *[("Dosen Ko-pembimbing" if lang == "id" else "Co-advisor", k) for k in META["ko_pembimbing"]],
    ]
    for label, sv in supervisors:
        center_para(doc, label, bold=True, size=12, space_before=6, space_after=0)
        center_para(doc, sv["nama"], size=12, space_after=0)
        center_para(doc, f"NIP {sv['nip']}", size=12, space_after=12)

    doc.add_paragraph()
    center_para(doc, prodi, size=12, space_before=12, space_after=0)
    center_para(doc, dept,  size=12, space_after=0)
    center_para(doc, fak,   size=12, space_after=0)
    center_para(doc, META["institusi"], size=12, space_after=0)
    center_para(doc, META["kota"], size=12, space_after=0)
    center_para(doc, META["tahun"], size=12, space_after=0)


def add_lembar_pengesahan(doc, lang="id"):
    judul_hal = "LEMBAR PENGESAHAN" if lang == "id" else "APPROVAL SHEET"
    judul_ta  = META["judul_id"] if lang == "id" else META["judul_en"]
    tipe_ta   = "TUGAS AKHIR" if lang == "id" else "FINAL PROJECT"
    kal1 = "Diajukan untuk memenuhi salah satu syarat" if lang == "id" else "Submitted to fulfill one of the requirements"
    kal2 = "memperoleh gelar Sarjana Komputer pada" if lang == "id" else "for obtaining a Bachelor of Computer Science degree at"
    oleh = "Oleh:" if lang == "id" else "By:"
    disetujui = "Disetujui oleh Tim Penguji Tugas Akhir:" if lang == "id" else "Approved by Final Project Examiner Team:"
    bulan = META["bulan"] if lang == "id" else META["bulan_en"]

    h = doc.add_heading(judul_hal, level=1)
    h.paragraph_format.page_break_before = True
    center_para(doc, judul_ta, bold=True, size=14, space_after=6)
    center_para(doc, tipe_ta, bold=True, size=14, space_after=12)
    center_para(doc, kal1, size=12, space_after=0)
    center_para(doc, kal2, size=12, space_after=0)
    center_para(doc, META["prodi_id"] if lang == "id" else META["prodi_en"], size=12, space_after=0)
    center_para(doc, META["dept_id"]  if lang == "id" else META["dept_en"],  size=12, space_after=0)
    center_para(doc, META["fak_id"]   if lang == "id" else META["fak_en"],   size=12, space_after=0)
    center_para(doc, META["institusi"], size=12, space_after=12)
    center_para(doc, f"{oleh} {META['nama']}", size=12, space_after=0)
    center_para(doc, f"NRP. {META['nrp']}", size=12, space_after=18)

    p = body_para(doc, disetujui, bold=True, indent=False)
    p.paragraph_format.space_after = Pt(6)

    all_sv = [
        (META["pembimbing"]["nama"], "Pembimbing" if lang == "id" else "Advisor"),
        *[(k["nama"], "Ko-pembimbing" if lang == "id" else "Co-advisor") for k in META["ko_pembimbing"]],
        *[(p_["nama"], "Penguji" if lang == "id" else "Examiner") for p_ in META["penguji"]],
    ]
    tbl = doc.add_table(rows=len(all_sv), cols=4)
    tbl.style = "Table Grid"
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = OxmlElement(f"w:{edge}")
                border.set(qn("w:val"), "none")
                tcBorders.append(border)
            tcPr.append(tcBorders)

    for i, (nama, jabatan) in enumerate(all_sv):
        cells = tbl.rows[i].cells
        cells[0].width = Cm(0.7)
        cells[1].width = Cm(6)
        cells[2].width = Cm(0.5)
        cells[3].width = Cm(6)

        def _set_cell(cell, text, bold=False):
            p = cell.paragraphs[0]
            p.style = doc.styles["Normal"]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(text)
            r.bold = bold

        _set_cell(cells[0], f"{i+1}.")
        # Nama + garis tanda tangan
        cell = cells[1]
        pc = cell.paragraphs[0]
        pc.style = doc.styles["Normal"]
        pc.paragraph_format.first_line_indent = Cm(0)
        pc.paragraph_format.space_before = Pt(36)  # ruang tanda tangan
        pc.paragraph_format.space_after  = Pt(0)
        pc.add_run(nama)
        _set_cell(cells[3], jabatan)

    doc.add_paragraph()
    center_para(doc, META["kota"].upper(), size=12, space_before=12, space_after=0)
    center_para(doc, f"{bulan}, {META['tahun']}", size=12, space_after=0)


def add_pernyataan_orisinalitas(doc, lang="id"):
    judul_hal = "PERNYATAAN ORISINALITAS" if lang == "id" else "STATEMENT OF ORIGINALITY"
    h = doc.add_heading(judul_hal, level=1)
    h.paragraph_format.page_break_before = True

    if lang == "id":
        body_para(doc, "Yang bertanda tangan di bawah ini:", indent=False)
    else:
        body_para(doc, "The undersigned:", indent=False)
    doc.add_paragraph()

    if lang == "id":
        rows = [
            ("Nama mahasiswa / NRP", ":", f"{META['nama']} / {META['nrp']}"),
            ("Program Studi", ":", "S-1 Teknik Informatika"),
            ("Dosen Pembimbing / NIP", ":", f"{META['pembimbing']['nama']} / {META['pembimbing']['nip']}"),
            ("Dosen Ko-pembimbing / NIP", ":", f"{META['ko_pembimbing'][0]['nama']} / {META['ko_pembimbing'][0]['nip']}"),
            ("", "", f"{META['ko_pembimbing'][1]['nama']} / {META['ko_pembimbing'][1]['nip']}"),
        ]
    else:
        rows = [
            ("Student Name / Student ID", ":", f"{META['nama']} / {META['nrp']}"),
            ("Study Program", ":", "Bachelor of Informatics"),
            ("Advisor / Employee ID", ":", f"{META['pembimbing']['nama']} / {META['pembimbing']['nip']}"),
            ("Co-advisor / Employee ID", ":", f"{META['ko_pembimbing'][0]['nama']} / {META['ko_pembimbing'][0]['nip']}"),
            ("", "", f"{META['ko_pembimbing'][1]['nama']} / {META['ko_pembimbing'][1]['nip']}"),
        ]
    add_2col_table(doc, rows, widths=(5, 0.4, 10))
    doc.add_paragraph()

    if lang == "id":
        body_para(doc, f'dengan ini menyatakan bahwa Tugas Akhir dengan judul “{META["judul_id"]}” adalah hasil karya sendiri, bersifat orisinal, dan ditulis dengan mengikuti kaidah penulisan ilmiah.', indent=False)
        doc.add_paragraph()
        body_para(doc, "Bilamana di kemudian hari ditemukan ketidaksesuaian dengan pernyataan ini, maka saya bersedia menerima sanksi sesuai dengan ketentuan yang berlaku di Institut Teknologi Sepuluh Nopember.", indent=False)
    else:
        body_para(doc, f'hereby declares that the Final Project entitled “{META["judul_en"]}” is my own work, is original, and was written in accordance with the rules of scientific writing.', indent=False)
        doc.add_paragraph()
        body_para(doc, "If any discrepancies with this statement are found in the future, I am willing to accept sanctions in accordance with the provisions of Institut Teknologi Sepuluh Nopember.", indent=False)

    doc.add_paragraph()
    # Tabel tanda tangan
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for edge in ("top","left","bottom","right","insideH","insideV"):
                border = OxmlElement(f"w:{edge}")
                border.set(qn("w:val"), "none")
                tcBorders.append(border)
            tcPr.append(tcBorders)

    c0, c1 = tbl.rows[0].cells
    def _sig_cell(cell, header, name, id_label, id_val):
        pHead = cell.paragraphs[0]
        pHead.style = doc.styles["Normal"]
        pHead.paragraph_format.first_line_indent = Cm(0)
        pHead.add_run(header).bold = True
        p2 = cell.add_paragraph()
        p2.style = doc.styles["Normal"]
        p2.paragraph_format.first_line_indent = Cm(0)
        p2.paragraph_format.space_before = Pt(42)
        p2.add_run(name)
        p3 = cell.add_paragraph()
        p3.style = doc.styles["Normal"]
        p3.paragraph_format.first_line_indent = Cm(0)
        p3.add_run(f"{id_label}. {id_val}")

    if lang == "id":
        _sig_cell(c0, "Mengetahui\nDosen Pembimbing",
                  META["pembimbing"]["nama"], "NIP", META["pembimbing"]["nip"])
        _sig_cell(c1, f"{META['kota']}, __________________\nMahasiswa",
                  META["nama"], "NRP", META["nrp"])
    else:
        _sig_cell(c0, "Acknowledged\nAdvisor",
                  META["pembimbing"]["nama"], "NIP", META["pembimbing"]["nip"])
        _sig_cell(c1, f"{META['kota']}, __________________\nStudent",
                  META["nama"], "NRP", META["nrp"])


def add_pernyataan_ai(doc):
    h = doc.add_heading("PERNYATAAN KODE ETIK PENGGUNAAN AI GENERATIF", level=1)
    h.paragraph_format.page_break_before = True
    p = body_para(doc, "Code of Conduct Statement: Generative AI or AI-Assisted Usage", indent=False)
    p.runs[0].italic = True
    doc.add_paragraph()

    body_para(doc, "Saya yang bertanda tangan di bawah ini: / I, the undersigned:", indent=False)
    rows = [
        ("Nama Mahasiswa / NRP\nFull Name / Student ID", ":", f"{META['nama']} / {META['nrp']}"),
        ("Program Studi / Study Program", ":", "S-1 Teknik Informatika"),
        ("Judul Tugas Akhir / Final Project Title", ":", META["judul_id"]),
    ]
    add_2col_table(doc, rows, widths=(5, 0.4, 10))
    doc.add_paragraph()
    body_para(doc, "dengan ini menyatakan bahwa pada Tugas Akhir dengan judul di atas tersebut:", indent=False)
    body_para(doc, "hereby declare that in the Final Project with the above title:", indent=False)
    doc.add_paragraph()

    tbl = doc.add_table(rows=3, cols=4)
    tbl.style = "Table Grid"
    headers = ["No.", "Pernyataan", "Statement", "Ya/Yes"]
    data = [
        ("1.", "Saya menggunakan AI Generatif dalam pengerjaan Tugas Akhir ini.",
         "I used Generative AI in this Final Project.", "[ ]"),
        ("2.", "Jika Ya, saya telah menyertakan lampiran penggunaan AI Generatif.",
         "If Yes, I have attached the Generative AI usage appendix.", "[ ]"),
    ]
    for j, hdr in enumerate(headers):
        p = tbl.rows[0].cells[j].paragraphs[0]
        p.style = doc.styles["Normal"]
        p.paragraph_format.first_line_indent = Cm(0)
        p.add_run(hdr).bold = True
    for i, row_d in enumerate(data):
        for j, val in enumerate(row_d):
            p = tbl.rows[i+1].cells[j].paragraphs[0]
            p.style = doc.styles["Normal"]
            p.paragraph_format.first_line_indent = Cm(0)
            p.add_run(val)

    doc.add_paragraph()
    p = body_para(doc, f"{META['kota']}, {META['bulan']} {META['tahun']}", indent=False)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_kata_pengantar(doc):
    h = doc.add_heading("KATA PENGANTAR", level=1)
    h.paragraph_format.page_break_before = True

    body_para(doc, f'Puji syukur penulis panjatkan kepada Allah SWT atas segala rahmat dan karunia-Nya sehingga penulis dapat menyelesaikan Tugas Akhir yang berjudul “{META["judul_id"]}” ini.')
    body_para(doc, "Penyusunan laporan Tugas Akhir ini tidak terlepas dari bantuan dan dukungan berbagai pihak. Oleh karena itu, penulis menyampaikan terima kasih kepada:")
    doc.add_paragraph()

    items = [
        f"{META['pembimbing']['nama']} selaku Dosen Pembimbing yang telah membimbing, memberikan masukan, dan motivasi kepada penulis.",
        f"{META['ko_pembimbing'][0]['nama']} selaku Dosen Ko-pembimbing yang telah memberikan bimbingan dan arahan.",
        f"{META['ko_pembimbing'][1]['nama']} selaku Dosen Ko-pembimbing yang telah memberikan bimbingan dan arahan.",
        "Seluruh dosen dan staf Departemen Teknik Informatika ITS.",
        "Keluarga dan teman-teman yang telah memberikan dukungan dan semangat.",
    ]
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(1)
        p.add_run(item)

    doc.add_paragraph()
    body_para(doc, "Penulis menyadari bahwa Tugas Akhir ini masih jauh dari sempurna. Oleh karena itu, penulis mengharapkan kritik dan saran yang membangun untuk penyempurnaan laporan ini.")
    doc.add_paragraph()
    p = body_para(doc, f"{META['kota']}, {META['bulan']} {META['tahun']}", indent=False)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()
    p = body_para(doc, "Penulis", indent=False)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_abstrak(doc, lang="id"):
    judul_hal = "ABSTRAK" if lang == "id" else "ABSTRACT"
    judul_ta  = META["judul_id"] if lang == "id" else META["judul_en"]
    h = doc.add_heading(judul_hal, level=1)
    h.paragraph_format.page_break_before = True

    center_para(doc, judul_ta, bold=True, size=12, space_after=12)

    if lang == "id":
        rows = [
            ("Nama Mahasiswa / NRP", ":", f"{META['nama']} / {META['nrp']}"),
            ("Departemen", ":", "Teknik Informatika FTEIC – ITS"),
            ("Dosen Pembimbing", ":", META["pembimbing"]["nama"]),
            ("Dosen Ko-pembimbing", ":", META["ko_pembimbing"][0]["nama"]),
            ("Dosen Ko-pembimbing", ":", META["ko_pembimbing"][1]["nama"]),
        ]
    else:
        rows = [
            ("Full Name / Student ID", ":", f"{META['nama']} / {META['nrp']}"),
            ("Department", ":", "Informatics ELECTICS – ITS"),
            ("Advisor", ":", META["pembimbing"]["nama"]),
            ("Co-advisor", ":", META["ko_pembimbing"][0]["nama"]),
            ("Co-advisor", ":", META["ko_pembimbing"][1]["nama"]),
        ]
    add_2col_table(doc, rows, widths=(4.5, 0.4, 10.5))
    doc.add_paragraph()

    if lang == "id":
        p = body_para(doc, "Abstrak", bold=True, indent=False)
        doc.add_paragraph()
        body_para(doc, "[ ISI ABSTRAK INDONESIA — tulis di sini ]")
        doc.add_paragraph()
        p = body_para(doc, f"Kata kunci: {META['kata_kunci_id']}", indent=False)
        p.runs[0].bold = True
    else:
        p = body_para(doc, "Abstract", bold=True, indent=False)
        doc.add_paragraph()
        body_para(doc, "[ ABSTRACT CONTENT IN ENGLISH — write here ]")
        doc.add_paragraph()
        p = body_para(doc, f"Keywords: {META['kata_kunci_en']}", indent=False)
        p.runs[0].bold = True


def add_daftar_isi_placeholder(doc):
    h = doc.add_heading("DAFTAR ISI", level=1)
    h.paragraph_format.page_break_before = True
    p = body_para(doc, "[ Daftar isi akan dibuat otomatis oleh Word: References → Table of Contents ]", indent=False)
    p.runs[0].italic = True

    h2 = doc.add_heading("DAFTAR GAMBAR", level=1)
    h2.paragraph_format.page_break_before = True
    p2 = body_para(doc, "[ Daftar gambar: References → Insert Table of Figures → pilih Figure ]", indent=False)
    p2.runs[0].italic = True

    h3 = doc.add_heading("DAFTAR TABEL", level=1)
    h3.paragraph_format.page_break_before = True
    p3 = body_para(doc, "[ Daftar tabel: References → Insert Table of Figures → pilih Table ]", indent=False)
    p3.runs[0].italic = True

    h4 = doc.add_heading("DAFTAR KODE SUMBER", level=1)
    h4.paragraph_format.page_break_before = True
    p4 = body_para(doc, "[ Daftar kode sumber — isi manual atau gunakan caption style khusus ]", indent=False)
    p4.runs[0].italic = True


def add_bab(doc, nomor, judul_id, subbab_list):
    h = doc.add_heading(f"BAB {nomor}\n{judul_id}", level=1)
    h.paragraph_format.page_break_before = True

    for sub in subbab_list:
        no_sub, judul_sub, sub3_list = sub
        doc.add_heading(f"{no_sub} {judul_sub}", level=2)
        body_para(doc, f"[ Isi sub-bab {no_sub} — {judul_sub} ]")
        for s3 in sub3_list:
            no3, judul3 = s3
            doc.add_heading(f"{no3} {judul3}", level=3)
            body_para(doc, f"[ Isi sub-sub-bab {no3} — {judul3} ]")


def add_daftar_pustaka(doc):
    h = doc.add_heading("DAFTAR PUSTAKA", level=1)
    h.paragraph_format.page_break_before = True
    body_para(doc, "[ Referensi ditulis dengan format IEEE. Contoh: ]", indent=False)
    doc.add_paragraph()

    refs = [
        '[1] Y. Zhu et al., "Author Recommendation System Based on Graph Neural Network," Information Sciences, vol. 562, pp. 1–15, 2021.',
        '[2] A. Mazharuddin Shiddiqi et al., "Findme-Scholar: Multi-Centroid Author Representation," Expert Systems with Applications, 2025.',
        '[3] A. Al Badawi et al., "OpenFHE: Open-Source Fully Homomorphic Encryption Library," in Proc. 10th Workshop on Encrypted Computing, 2022, pp. 28–39.',
    ]
    for ref in refs:
        p = body_para(doc, ref, indent=False)
        p.paragraph_format.left_indent  = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(-1.25)  # hanging indent


def add_lampiran(doc):
    h = doc.add_heading("LAMPIRAN", level=1)
    h.paragraph_format.page_break_before = True
    doc.add_heading("Lampiran A — Kode Sumber Lengkap", level=2)
    body_para(doc, "[ Cantumkan atau referensikan kode sumber di sini. ]")
    doc.add_heading("Lampiran B — Log Penggunaan AI Generatif", level=2)
    body_para(doc, "[ Cantumkan log percakapan / prompt yang digunakan dengan AI Generatif. ]")
    doc.add_heading("Lampiran C — Data Pengujian", level=2)
    body_para(doc, "[ Cantumkan data-data pengujian jika diperlukan. ]")


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────

def build():
    doc = Document()
    set_margins(doc)
    add_styles(doc)

    # ── SAMPUL ─────────────────────────────────
    # Hapus paragraf default kosong pertama
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    add_cover(doc, lang="id")
    add_blank_intentional(doc)
    add_cover(doc, lang="en")
    add_blank_intentional(doc)

    # ── FRONT MATTER ───────────────────────────
    add_lembar_pengesahan(doc, lang="id")
    add_blank_intentional(doc)
    add_lembar_pengesahan(doc, lang="en")
    add_blank_intentional(doc)

    add_pernyataan_orisinalitas(doc, lang="id")
    add_blank_intentional(doc)
    add_pernyataan_orisinalitas(doc, lang="en")
    add_blank_intentional(doc)

    add_pernyataan_ai(doc)
    add_blank_intentional(doc)

    add_kata_pengantar(doc)
    add_blank_intentional(doc)

    add_abstrak(doc, lang="id")
    add_blank_intentional(doc)
    add_abstrak(doc, lang="en")
    add_blank_intentional(doc)

    add_daftar_isi_placeholder(doc)

    # ── MAIN MATTER ────────────────────────────
    add_bab(doc, "1", "PENDAHULUAN", [
        ("1.1", "Latar Belakang", []),
        ("1.2", "Rumusan Masalah", []),
        ("1.3", "Batasan Masalah", []),
        ("1.4", "Tujuan", []),
        ("1.5", "Manfaat", [
            ("1.5.1", "Manfaat Teoritis"),
            ("1.5.2", "Manfaat Praktis"),
            ("1.5.3", "Manfaat Sosial"),
        ]),
    ])

    add_bab(doc, "2", "TINJAUAN PUSTAKA", [
        ("2.1", "Hasil Penelitian Terdahulu", []),
        ("2.2", "Dasar Teori", [
            ("2.2.1", "Sistem Rekomendasi"),
            ("2.2.2", "Author Recommendation System"),
            ("2.2.3", "Embedding Semantik untuk Representasi Author"),
            ("2.2.4", "K-Means Clustering"),
            ("2.2.5", "Cosine Similarity"),
            ("2.2.6", "Privasi pada Recommendation System"),
            ("2.2.7", "Homomorphic Encryption"),
            ("2.2.8", "Homomorphic Inference"),
        ]),
    ])

    add_bab(doc, "3", "METODOLOGI", [
        ("3.1", "Metode yang Digunakan", [
            ("3.1.1", "Arsitektur Sistem Rekomendasi Author"),
            ("3.1.2", "Pembentukan Representasi Author Berbasis Multi-Centroid"),
            ("3.1.3", "Pemrosesan Query Pengguna di Sisi Client"),
            ("3.1.4", "Enkripsi Query dan Homomorphic Similarity Computation"),
            ("3.1.5", "Dekripsi dan Pemeringkatan Hasil Rekomendasi"),
        ]),
        ("3.2", "Bahan dan Peralatan yang Digunakan", []),
        ("3.3", "Urutan Pelaksanaan Penelitian", []),
    ])

    add_bab(doc, "4", "PERANCANGAN DAN IMPLEMENTASI", [
        ("4.1", "Perancangan Sistem", [
            ("4.1.1", "Perancangan Struktur Data"),
            ("4.1.2", "Perancangan Antarmuka"),
        ]),
        ("4.2", "Implementasi", [
            ("4.2.1", "Implementasi Pembentukan Representasi Author"),
            ("4.2.2", "Implementasi Enkripsi Query dengan FHE"),
            ("4.2.3", "Implementasi Homomorphic Similarity Computation"),
            ("4.2.4", "Implementasi Dekripsi dan Ranking"),
        ]),
    ])

    add_bab(doc, "5", "PENGUJIAN DAN EVALUASI", [
        ("5.1", "Lingkungan Pengujian", []),
        ("5.2", "Skenario Pengujian", [
            ("5.2.1", "Pengujian Fungsional"),
            ("5.2.2", "Pengujian Akurasi Rekomendasi"),
            ("5.2.3", "Pengujian Performa (Waktu Komputasi)"),
        ]),
        ("5.3", "Analisis Hasil", []),
    ])

    add_bab(doc, "6", "PENUTUP", [
        ("6.1", "Kesimpulan", []),
        ("6.2", "Saran", []),
    ])

    add_daftar_pustaka(doc)
    add_lampiran(doc)

    out = "TA-Ricardo-template.docx"
    doc.save(out)
    print(f"[OK] File berhasil dibuat: {out}")


if __name__ == "__main__":
    build()
